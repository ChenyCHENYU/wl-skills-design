#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
从原版参考文档提取骨架：
  保留 Section 0（封面 1/封面 2/修订记录/目录）
  剔除 Section 1 的全部章节内容
  输出 template_skeleton.docx → 作为仓库里的格式基准

封面中的元数据（项目名、卷名、公司名等）保持原始值，
generate_spec_doc.py 在加载骨架后做文字替换更新。
"""
import os, sys
from docx import Document
from docx.oxml.ns import qn

_HERE     = os.path.dirname(os.path.abspath(__file__))   # spec-gen/scripts/
_SPEC_GEN = os.path.dirname(_HERE)                       # spec-gen/

TPLPATH = r'D:\work\【15】日常工作\【#】产品设计\烟台华新数智化改造项目一期_需求设计说明书_系统需求设计_生产管理分册.docx'
# ⚠️  骨架模板写入 templates/ — 生成后作为只读基准，不要再次运行覆盖
OUT     = os.path.join(_SPEC_GEN, 'templates', 'template_skeleton.docx')

def extract_skeleton(src: str, dst: str):
    if not os.path.exists(src):
        print(f'❌  原始模板不存在：{src}')
        sys.exit(1)

    doc  = Document(src)
    body = doc.element.body

    # ── 策略：找最后一个 toc 样式段落，在它之后截断 ─────────────────────
    # 原版结构：Section0(封面) → Section1(修订记录 + 目录条目 + 章节正文)
    # 正确骨架 = 封面 + 修订记录 + 目录条目，不含任何章节正文
    # toc 1/toc 2/toc 3 是目录条目专用样式，最后一条之后即为章节正文开始
    last_toc_p = None
    for para in doc.paragraphs:
        sname = para.style.name.lower()
        if sname.startswith('toc') or sname.startswith('目录'):
            last_toc_p = para._p

    # 备用：如果没有 toc 样式，找第二个 Heading 1（修订记录=第1个，章节=第2个）
    if last_toc_p is None:
        print('⚠️  未找到 toc 样式段落，改用 Heading 1 定位')
        h1_seen = 0
        for para in doc.paragraphs:
            if para.style.name == 'Heading 1':
                h1_seen += 1
                if h1_seen == 2:          # 第 2 个 H1 = 第一章，在它前面截断
                    siblings = list(body)
                    idx = siblings.index(para._p)
                    last_toc_p = siblings[idx - 1] if idx > 0 else None
                    break

    if last_toc_p is None:
        print('❌  无法确定切割点，退出')
        return

    # ── 删除 last_toc_p 之后的所有内容（保留最终 sectPr）────────────────
    found  = False
    to_del = []
    for child in list(body):
        if child is last_toc_p:
            found = True
            continue            # 保留最后一个 toc 段落本身
        if found:
            if child.tag == qn('w:sectPr'):   # 最终 sectPr 必须保留
                continue
            to_del.append(child)

    removed = len(to_del)
    for elem in to_del:
        body.remove(elem)

    doc.save(dst)

    # ── 验证输出 ─────────────────────────────────────────────────────────
    v = Document(dst)
    print(f'✅  骨架已输出：{dst}')
    print(f'   共剔除 {removed} 个元素（章节正文）')
    print(f'   保留段落数：{len(v.paragraphs)}  表格数：{len(v.tables)}  节数：{len(v.sections)}')
    print()
    print('骨架结构预览：')
    for i, p in enumerate(v.paragraphs):
        txt = p.text.strip()
        if txt:
            print(f'  [{i:02d}] {p.style.name:20} | {txt[:55]}')
    print()
    print('骨架表格：')
    for i, t in enumerate(v.tables):
        r0 = [c.text[:12] for c in t.rows[0].cells] if t.rows else []
        print(f'  表{i:02d} ({len(t.rows)}行×{len(t.columns)}列): {r0}')
    print()
    print('下一步：')
    print('  1. 用 WPS 打开 template_skeleton.docx 确认封面/修订/目录完整')
    print('  2. git add template_skeleton.docx 提交到仓库')
    print('  3. generate_spec_doc.py 改为 Document("template_skeleton.docx") 注入章节')


if __name__ == '__main__':
    extract_skeleton(TPLPATH, OUT)
