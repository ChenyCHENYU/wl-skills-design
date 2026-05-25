#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成 华新计划模块需求设计说明书 Word 文档
精准还原金恒信息科技标准 Word 模板格式：
  - 页眉：右对齐公司名（中英文）+ 横线
  - 页脚：3 行元数据表格
  - 封面 1（信息封面）/ 封面 2（签字页）
  - 修订记录 / 自动目录域 / 正文章节
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, re, os, sys

# ── 目录锚点 ──────────────────────────────────────────────────
_HERE     = os.path.dirname(os.path.abspath(__file__))   # spec-gen/scripts/
_SPEC_GEN = os.path.dirname(_HERE)                       # spec-gen/
sys.path.insert(0, _HERE)   # 确保能 import 同目录的 draw_flow

from draw_flow import (generate_flow_png, generate_flow_drawio,
                        img_placeholder_png, FLOW_PMPM_A01, FLOW_PMPM_A06)

# ─────────────────────────────────────────────────────────────
#  项目元数据
# ─────────────────────────────────────────────────────────────
META = {
    'client':      '烟台华新不锈钢有限公司',
    'project':     '烟台华新数智化改造项目一期',
    'vol':         '第一卷  计划模块',
    'proj_no':     'HXSS-2026-001',
    'doc_no':      'HX-SRS-PM-001',
    'version':     '1.0',
    'author':      '高钰',
    'date':        '2026.3.25',
    'company':     '江苏金恒信息科技股份有限公司',
    'company_en':  'Jiangsu Jinheng Information Technology Co.,Ltd.',
    'doc_name':    '需求设计说明书_计划模块_系统需求设计',
}

FONT_CN   = '宋体'
FONT_HEAD = '黑体'
FONT_EN   = 'Times New Roman'

# 参考模板路径（优先使用仓库内骨架模板，外部原版作为备选）
_TPL_SKELETON = os.path.join(_SPEC_GEN, 'templates', 'template_skeleton.docx')
_TPL_ORIGINAL = r'D:\work\【15】日常工作\【#】产品设计\烟台华新数智化改造项目一期_需求设计说明书_系统需求设计_生产管理分册.docx'
TPLPATH = _TPL_SKELETON if os.path.exists(_TPL_SKELETON) else _TPL_ORIGINAL

GRAY_LIGHT = 'D9D9D9'   # 表头背景色
GRAY_MID   = 'BFBFBF'

# ─────────────────────────────────────────────────────────────
#  XML 辅助函数
# ─────────────────────────────────────────────────────────────
def _rpr(run):
    """确保 run 有 rPr 节点，返回它"""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    return rPr

def set_run_east_asia(run, font_name):
    rPr = _rpr(run)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def set_run(run, cn_font=FONT_CN, size=10.5, bold=False, italic=False, color=None):
    run.font.name = cn_font
    set_run_east_asia(run, cn_font)
    run.font.size    = Pt(size)
    run.font.bold    = bold
    run.font.italic  = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_fmt(para, align=None, sb=0, sa=6, ls=None):
    pf = para.paragraph_format
    if align is not None: pf.alignment     = align
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    if ls is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = ls

def add_para_bottom_border(para, sz=6, color='000000'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(sz))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_cell_bg(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def set_cell_valign(cell, align='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def set_table_borders(table, color='000000', sz=4):
    tbl    = table._tbl
    tblPr  = tbl.tblPr
    tblBdr = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBdr.append(el)
    tblPr.append(tblBdr)

def set_table_width(table, width_cm):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    tblW  = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(width_cm * 567)))  # EMU/twips
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def cell_write(cell, text, cn_font=FONT_CN, size=10, bold=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, bg=None, valign='center'):
    para = cell.paragraphs[0]
    # clear existing runs
    for run in para.runs:
        run._r.getparent().remove(run._r)
    # 支持 \n 软换行（在同一单元格段落内插入 <w:br>）
    lines = str(text).split('\n')
    for idx, line in enumerate(lines):
        run = para.add_run(line)
        set_run(run, cn_font=cn_font, size=size, bold=bold)
        if idx < len(lines) - 1:
            br = OxmlElement('w:br')
            run._r.append(br)
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)
    if bg:
        set_cell_bg(cell, bg)
    set_cell_valign(cell, valign)
    return para

def page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

def insert_page_field(para, center=False):
    """在段落中插入 {PAGE} 域"""
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    for tag, ftype in (('w:fldChar','begin'),('w:fldChar','end')):
        pass  # placeholder
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    it  = OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text = ' PAGE '
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fc1, it, fc2, fc3])

def blank_para(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        para_fmt(p, sb=0, sa=0)

# ─────────────────────────────────────────────────────────────
#  页眉 / 页脚
# ─────────────────────────────────────────────────────────────
def setup_header(section):
    hdr = section.header
    # 清空默认段落内容
    for p in hdr.paragraphs:
        for r in p.runs:
            r._r.getparent().remove(r._r)

    # 第一行：中文公司名，右对齐
    p1 = hdr.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(META['company'])
    set_run(r1, cn_font=FONT_CN, size=10.5, bold=False)
    para_fmt(p1, sb=0, sa=1)

    # 第二行：英文公司名，右对齐，带下边框（横线）
    p2 = hdr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(META['company_en'])
    r2.font.name = FONT_EN
    set_run_east_asia(r2, FONT_EN)
    r2.font.size = Pt(9)
    para_fmt(p2, sb=0, sa=2)
    add_para_bottom_border(p2, sz=8, color='1F3864')  # 深蓝色横线

def setup_footer(section):
    ftr = section.footer
    # 保留最后一个段落（Word 要求 footer 末尾必须有 <w:p>）
    # 先把 table 插入到默认段落之前
    tbl = ftr.add_table(rows=3, cols=6, width=Inches(6.3))
    # 将 table 移到 footer 首位（在默认 <w:p> 前面）
    tbl_elem  = tbl._tbl
    first_p   = ftr.paragraphs[0]._p if ftr.paragraphs else None
    if first_p is not None:
        ftr._element.remove(tbl_elem)
        ftr._element.insert(list(ftr._element).index(first_p), tbl_elem)
    # 让尾段落无高度
    if ftr.paragraphs:
        pf = ftr.paragraphs[-1].paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color='595959', sz=4)

    # 列宽（单位 cm，合计约 16.5 cm 对应 A4 可用宽）
    col_w = [Cm(1.8), Cm(4.2), Cm(2.0), Cm(5.0), Cm(1.8), Cm(2.2)]
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(int(col_w[j].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # ── 第 1 行 ──
    r0 = tbl.rows[0]
    cell_write(r0.cells[0], '项  目  号', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r0.cells[1], META['proj_no'], size=8)
    cell_write(r0.cells[2], '文档编号', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r0.cells[3], META['doc_no'], size=8)
    cell_write(r0.cells[4], '编制日期', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r0.cells[5], META['date'], size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 第 2 行 ──
    r1 = tbl.rows[1]
    cell_write(r1.cells[0], '项目名称', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r1.cells[1], META['project'], size=8)
    cell_write(r1.cells[2], '文档名称', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r1.cells[3], META['doc_name'], size=8)
    cell_write(r1.cells[4], '设计专业', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r1.cells[5], '应用软件', size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 第 3 行 ──
    r2 = tbl.rows[2]
    cell_write(r2.cells[0], '项目阶段', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r2.cells[1], '需求确认与详细设计', size=8)
    cell_write(r2.cells[2], '页  号', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    # 页码域
    r2.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = r2.cells[3].paragraphs[0].paragraph_format
    pf.space_before = Pt(1); pf.space_after = Pt(1)
    insert_page_field(r2.cells[3].paragraphs[0], center=True)
    cell_write(r2.cells[4], '保密级别', size=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
    cell_write(r2.cells[5], '企业秘密', size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER)

# ─────────────────────────────────────────────────────────────
#  封面 1 — 信息封面
# ─────────────────────────────────────────────────────────────
def build_cover1(doc):
    """封面 — 严格匹配参考模板：居中 / 黑体 / 精确字号"""
    blank_para(doc, 4)

    def ctitle(text, size, bold=False, sb=0, sa=12):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run(r, cn_font=FONT_HEAD, size=size, bold=bold)
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=sb, sa=sa)
        return p

    # 模板封面层级： project=26pt / title=36pt / vol=24pt / subtitle=26pt
    ctitle(META['project'], 26, bold=True,  sb=0, sa=12)
    ctitle('需求设计说明书', 36, bold=True,  sb=0, sa=18)
    ctitle(META['vol'],     24, bold=False, sb=0, sa=12)
    ctitle('系统需求设计', 26, bold=False, sb=0, sa=20)

    # 信息表（2 列）
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    col_w = [Cm(4), Cm(9)]
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_w[j].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    rows_data = [
        ('项  目  号', META['proj_no']),
        ('文 档 编 号', META['doc_no']),
        ('版  本  号',  META['version']),
        ('保  密  级  别', '一般  □    内部公开  □    企业秘密  ☑    企业机密  □'),
    ]
    for i, (k, v) in enumerate(rows_data):
        cell_write(tbl.rows[i].cells[0], k, size=11, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
        cell_write(tbl.rows[i].cells[1], v, size=11)

    blank_para(doc, 3)

    p = doc.add_paragraph()
    r = p.add_run(META['company'])
    set_run(r, cn_font=FONT_HEAD, size=18, bold=True)
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)

    page_break(doc)

# ─────────────────────────────────────────────────────────────
#  封面 2 — 签字页
# ─────────────────────────────────────────────────────────────
def build_cover2(doc):
    blank_para(doc, 5)

    def ctr(text, font, size, bold=False, sb=6, sa=6):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run(r, cn_font=font, size=size, bold=bold)
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=sb, sa=sa)
        return p

    ctr(META['client'],  FONT_HEAD, 18, bold=True, sb=0, sa=8)
    ctr(META['project'], FONT_HEAD, 16, bold=True, sb=0, sa=20)
    ctr('需求设计说明书', FONT_HEAD, 26, bold=True, sb=0, sa=14)
    ctr(META['vol'],     FONT_HEAD, 18, bold=False, sb=0, sa=8)
    ctr('系统需求设计',  FONT_HEAD, 14, bold=False, sb=0, sa=30)

    # 签字表（4 列）
    tbl = doc.add_table(rows=3, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    col_w = [Cm(2.5), Cm(4.5), Cm(2.5), Cm(4.5)]
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_w[j].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    rows_data = [
        ('编  制', META['author'],  '编制日期', META['date']),
        ('审  核', 'XX',            '审核日期', '20XX.XX.XX'),
        ('批  准', 'XX',            '批准日期', '20XX.XX.XX'),
    ]
    for i, (k1, v1, k2, v2) in enumerate(rows_data):
        cell_write(tbl.rows[i].cells[0], k1, size=11, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
        cell_write(tbl.rows[i].cells[1], v1, size=11)
        cell_write(tbl.rows[i].cells[2], k2, size=11, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)
        cell_write(tbl.rows[i].cells[3], v2, size=11)

    blank_para(doc, 5)

    p = doc.add_paragraph()
    r = p.add_run(META['company'])
    set_run(r, cn_font=FONT_HEAD, size=14)
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)

    page_break(doc)

# ─────────────────────────────────────────────────────────────
#  修订记录页
# ─────────────────────────────────────────────────────────────
def build_revision(doc):
    # 标题（使用 Heading 1 样式，保证目录能识别）
    try:
        p = doc.add_paragraph(style='Heading 1')
    except Exception:
        p = doc.add_paragraph()
    for run in list(p.runs):
        run._r.getparent().remove(run._r)
    r = p.add_run('修  订  记  录')
    set_run(r, cn_font=FONT_HEAD, size=16, bold=True)
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=24, sa=18)

    headers = ['日期', '作者', '版本号', '修改原因', '主要修改内容']
    widths  = [Cm(2.5), Cm(1.8), Cm(1.8), Cm(3.5), Cm(7.4)]

    tbl = doc.add_table(rows=9, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    for j, w in enumerate(widths):
        for row in tbl.rows:
            tc = row.cells[j]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w.pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # 表头行
    for j, h in enumerate(headers):
        cell_write(tbl.rows[0].cells[j], h, cn_font=FONT_HEAD,
                   size=10.5, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=GRAY_LIGHT)

    # 第一数据行
    row1_data = [META['date'], META['author'], '1.0', '新建', '/']
    for j, v in enumerate(row1_data):
        cell_write(tbl.rows[1].cells[j], v, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 空行（高度撑开）
    for i in range(2, 9):
        for j in range(5):
            cell_write(tbl.rows[i].cells[j], '', size=10)
            # 设置行高
        tbl.rows[i].height = Cm(0.8)

    page_break(doc)

# ─────────────────────────────────────────────────────────────
#  目录页
# ─────────────────────────────────────────────────────────────
def build_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run('目  录')
    set_run(r, cn_font=FONT_HEAD, size=16, bold=True)
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=24, sa=18)

    # 插入 Word 自动目录域（文档打开时由 updateFields 自动更新）
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    it  = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = ' TOC \\o "1-4" \\h \\z \\u '
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
    run._r.extend([fc1, it, fc2])

    # 预填目录条目（updateFields 会替换这些内容）
    toc_entries = [
        (1, '修订记录', ''),
        (1, '1.  系统目标', ''),
        (1, '2.  组织架构及业务单位', ''),
        (2, '2.1  现有组织架构', ''),
        (2, '2.2  岗位定义', ''),
        (2, '2.3  专有名称', ''),
        (1, '3.  总体设计', ''),
        (2, '3.1  业务架构', ''),
        (2, '3.2  系统流程', ''),
        (3, '3.2.1  系统流程总图', ''),
        (3, '3.2.2  系统流程清单', ''),
        (2, '3.3  功能层级', ''),
        (1, '4.  详细设计', ''),
        (2, '4.3  计划管理', ''),
        (3, '4.3.1  系统流程清单', ''),
        (3, '4.3.2  系统流程说明', ''),
        (4, '4.3.2.1  PMPM-A-01  炼钢计划编制流程', ''),
        (3, '4.3.3  系统流程与作业画面对照表', ''),
        (3, '4.3.4  功能设计', ''),
        (4, '4.3.4.7  【PMPM007】炼钢计划编制', ''),
        (3, '4.3.5  内部接口（API）', ''),
    ]
    # 目录条目尺寸参考原文档：
    #   H1 章目: 12pt bold  H2 节目: 11pt  H3 小节: 10.5pt  H4 功能: 10pt
    toc_size = {1: 12, 2: 11, 3: 10.5, 4: 10}
    toc_bold = {1: True, 2: False, 3: False, 4: False}
    indent_map = {1: Pt(0), 2: Pt(18), 3: Pt(36), 4: Pt(54)}
    for level, title, page in toc_entries:
        ep = doc.add_paragraph()
        ep.paragraph_format.left_indent = indent_map.get(level, Pt(36))
        ep.paragraph_format.space_before = Pt(1.5)
        ep.paragraph_format.space_after  = Pt(1.5)
        er = ep.add_run(title)
        er.font.name = FONT_CN
        set_run_east_asia(er, FONT_CN)
        er.font.size = Pt(toc_size.get(level, 10))
        er.font.bold = toc_bold.get(level, False)

    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')
    # 把 fc3 加到最后一个条目段落的最后一个 run 上
    last_run = doc.paragraphs[-1].add_run()
    last_run._r.append(fc3)

    page_break(doc)

# ─────────────────────────────────────────────────────────────
#  正文辅助：标题 / 正文 / 图片占位 / 列表
# ─────────────────────────────────────────────────────────────
def h(doc, text, level):
    """生成章节标题，level 1-4
    参考模板实测字号（均继承模板样式内 SimSun 默认字体）：
      H1 = 宋体 14pt Bold  （章级，样式内 pageBreakBefore）
      H2 = 宋体 14pt Bold  （节级）
      H3 = 宋体 14pt Regular（小节级）
      H4 = 宋体 14pt Regular（功能级）
    """
    styles = {
        1: ('Heading 1', FONT_CN, 14, True,  18, 6),
        2: ('Heading 2', FONT_CN, 14, True,  12, 6),
        3: ('Heading 3', FONT_CN, 14, False,  9, 6),
        4: ('Heading 4', FONT_CN, 14, False,  6, 4),
        5: ('Heading 5', FONT_CN, 12, False,  4, 3),
    }
    sname, font, size, bold, sb, sa = styles.get(level, styles[4])
    try:
        p = doc.add_paragraph(style=sname)
    except Exception:
        p = doc.add_paragraph()
    # 清空再写，避免双内容
    for run in list(p.runs):
        run._r.getparent().remove(run._r)
    r = p.add_run(text)
    set_run(r, cn_font=font, size=size, bold=bold)
    para_fmt(p, sb=sb, sa=sa)
    return p

def body(doc, text, first_indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, cn_font=FONT_CN, size=10.5)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.5
    if first_indent:
        pf.first_line_indent = Pt(21)  # 两字符缩进
    return p

def img_placeholder(doc, caption):
    p = doc.add_paragraph()
    r = p.add_run(f'【此处插入：{caption}】')
    set_run(r, cn_font=FONT_CN, size=10.5, italic=True, color=(127,127,127))
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=6)
    # 虚线框效果（段落边框）
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'dashed')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '7F7F7F')
        pBdr.append(el)
    pPr.append(pBdr)
    return p

# ─── 生成的图片临时目录 ─────────────────────────────────────────────
# 输出目录（所有生成物写到 spec-gen/output/，不污染模板）
_OUT_DIR = os.path.join(_SPEC_GEN, 'output')
_IMG_DIR = os.path.join(_OUT_DIR, 'assets')
os.makedirs(_IMG_DIR, exist_ok=True)
os.makedirs(_OUT_DIR, exist_ok=True)

def insert_flow_img(doc, png_path, width_cm=16.0, caption=None):
    """将 PNG 插入文档，居中，可选图注"""
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(png_path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        set_run(r, cn_font=FONT_CN, size=9, color=(89, 89, 89), italic=True)
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(6)
    return p

def note(doc, text):
    """说明/注释段落（带缩进，浅色）"""
    p = doc.add_paragraph()
    r = p.add_run(f'说明：{text}')
    set_run(r, cn_font=FONT_CN, size=10, italic=False, color=(89,89,89))
    pf = p.paragraph_format
    pf.left_indent  = Pt(21)
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    return p

# ─────────────────────────────────────────────────────────────
#  通用数据表生成器
# ─────────────────────────────────────────────────────────────
def make_table(doc, headers, rows_data, col_widths_cm, header_bg=GRAY_LIGHT, font_size=10):
    """
    headers       : list[str]
    rows_data     : list[list[str]]
    col_widths_cm : list[float]
    """
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)
    set_table_width(tbl, sum(col_widths_cm))

    # 设置列宽
    for row in tbl.rows:
        for j, w in enumerate(col_widths_cm):
            tc = row.cells[j]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(Cm(w).pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # 表头
    for j, h_text in enumerate(headers):
        cell_write(tbl.rows[0].cells[j], h_text,
                   cn_font=FONT_HEAD, size=font_size, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=header_bg)

    # 数据行
    for i, row_data in enumerate(rows_data):
        for j, cell_txt in enumerate(row_data):
            cell_write(tbl.rows[i+1].cells[j], str(cell_txt),
                       cn_font=FONT_CN, size=font_size,
                       align=WD_ALIGN_PARAGRAPH.LEFT)

    # 表后间距
    p = doc.add_paragraph()
    para_fmt(p, sb=0, sa=6)
    return tbl

# ─────────────────────────────────────────────────────────────
#  第 1 章：系统目标
# ─────────────────────────────────────────────────────────────
def build_ch1(doc):
    h(doc, '1.  系统目标', 1)

    goals = [
        '适应少量多样的市场需求，以生产订单组织生产，通过订单交期评估实现对外承诺交期的精准管理。',
        '销售、质量、物流、装备等业务模块协同，基于系统建立横向贯通、纵向到底的一贯生产管理体系。',
        '按计划管控生产节奏，追踪生产工单执行状态，实时掌握生产进度与异常情况。',
        '对全流程生产数据建立多维分析模型，为各级技术及管理人员提供决策支持。',
    ]
    for i, goal in enumerate(goals, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. {goal}')
        set_run(r, cn_font=FONT_CN, size=10.5)
        pf = p.paragraph_format
        pf.left_indent    = Pt(21)
        pf.space_before   = Pt(0)
        pf.space_after    = Pt(8)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = 1.5

# ─────────────────────────────────────────────────────────────
#  第 2 章：组织架构及业务单位
# ─────────────────────────────────────────────────────────────
def build_ch2(doc):
    h(doc, '2.  组织架构及业务单位', 1)

    # 2.1
    h(doc, '2.1  现有组织架构', 2)
    img_placeholder(doc, '烟台华新不锈钢有限公司 组织架构图')
    note(doc, '公司采用三级管控体系——集团管控层（总经理室）、工厂经营层（各职能部门）、'
              '车间执行层（各生产工序）。生产管理部门下设计划科、调度科、技术科，'
              '分别承担生产计划编制、生产调度指挥和工艺技术管理职责。')

    # 2.2 岗位定义
    h(doc, '2.2  岗位定义', 2)
    hdrs = ['序号', '岗位名称', '所属部门', '岗位职责（摘要）', '在本系统中的角色']
    data = [
        ['1', '计划员', '生产管理部-计划科',
         '负责月度、周、日生产计划的编制与下达，协调各工序生产节奏',
         '炼钢计划新增、修改、确认、下发；利库计划编制；计划查询统计'],
        ['2', '调度员', '生产管理部-调度科',
         '负责生产过程中的实时调度指挥，处理生产异常',
         '查看已下发计划；生产工单跟踪；异常上报'],
        ['3', '生产主管', '生产管理部',
         '负责整体生产计划审批及资源调配',
         '计划审批；计划完成率看板查看；多维度统计分析'],
        ['4', '炼钢工长', '炼钢车间',
         '负责炼钢工序的生产执行，按计划组织班组生产',
         '查看炼钢计划（只读）；实绩录入'],
        ['5', '技术员', '生产管理部-技术科',
         '负责钢种工艺参数维护及技术标准更新',
         '钢种合金配比维护；工艺参数查询'],
        ['6', '物料员', '物料部',
         '负责原料采购计划和库存管理',
         'BOM原料需求计算结果查看；利库申请处理'],
    ]
    make_table(doc, hdrs, data, [0.8, 2.0, 3.0, 5.0, 5.2])

    # 2.3 专有名称
    h(doc, '2.3  专有名称', 2)
    hdrs2 = ['序号', '术语 / 专有名称', '类别', '定义说明']
    data2 = [
        ['1', '炼钢序号', '业务术语',
         '同一炼钢订单下，按工序顺序编排的序号，格式为纯数字，从1开始递增；相同炼钢订单下炼钢序号不可重复'],
        ['2', '炼钢订单', '业务术语',
         '由订单管理模块生成、驱动炼钢工序生产的工作指令，包含钢种、规格、数量、交期等关键信息；是炼钢计划的直接输入依据'],
        ['3', '利库', '业务术语',
         '使用仓库现有库存来满足生产需求，而不新购物料的行为；本系统支持利库计划的编制和审批'],
        ['4', 'BIP', '外部系统术语',
         '金恒信息自研的 ERP 平台；本系统从 BIP 同步销售订单、物料数据、BOM 数据等基础数据'],
        ['5', '钢种', '业务术语',
         '按化学成分和力学性能分类的钢材品种标识，如 304、316L 等；是生产计划的核心分类维度，与合金配比一一对应'],
        ['6', '轧线', '业务术语',
         '钢材热轧或冷轧的生产线，烟台华新现有多条轧线；炼钢计划需指定所属轧线，以便调度员按线排产'],
        ['7', '合炉率', '业务术语',
         '同一炉次中，相同钢种订单合并冶炼的比例；用于评估生产效率，值越高越节省冶炼成本'],
    ]
    make_table(doc, hdrs2, data2, [0.8, 2.5, 2.5, 10.2])

# ─────────────────────────────────────────────────────────────
#  第 3 章：总体设计
# ─────────────────────────────────────────────────────────────
def build_ch3(doc):
    h(doc, '3.  总体设计', 1)

    # 3.1
    h(doc, '3.1  业务架构', 2)
    img_placeholder(doc, '计划模块 业务架构图')
    note(doc, '计划模块业务架构分为三层。外部系统对接层：与 BIP（ERP）双向对接，从 BIP 同步销售订单、'
              'BOM 物料数据、库存数据作为计划输入，向 BIP 回传生产计划确认结果；与生产执行 MES 对接，'
              '下发炼钢计划、利库计划。核心功能模块层：包含目标管理、订单管理、计划管理三个子模块，'
              '形成「目标→订单→计划」的纵向业务链条。基础数据层：钢种合金配比、工作天数、'
              '生产日历、轧线产能等基础参数，为计划编制提供约束依据。')

    # 3.2
    h(doc, '3.2  系统流程', 2)

    h(doc, '3.2.1  系统流程总图', 3)
    img_placeholder(doc, '计划模块 系统流程总图')
    note(doc, '泳道分配：纵向按岗位划分泳道——计划员、调度员、生产主管、炼钢工长、物料员。'
              '流程起点为「收到销售订单」，终点为「计划下发至MES并开始生产执行」。')

    h(doc, '3.2.2  系统流程清单', 3)
    hdrs = ['子模块', '流程编码', '流程名称', '流程简述']
    data = [
        ['目标管理', 'PMMB-A-01', '年度生产目标流程', '年初制定年度生产目标并按钢种、轧线分解到月度'],
        ['目标管理', 'PMMB-A-02', '月度生产目标流程', '每月根据年度目标和订单情况制定月度生产目标'],
        ['订单管理', 'PMOM-A-01', '炼钢订单导入流程', '从BIP同步销售订单并生成炼钢订单'],
        ['订单管理', 'PMOM-A-02', '炼钢订单交期评估流程', '对炼钢订单进行产能评估并确认交期'],
        ['计划管理', 'PMPM-A-01', '炼钢计划编制流程', '计划员编制炼钢计划，经审核后下发至炼钢工序'],
        ['计划管理', 'PMPM-A-02', '利库计划流程', '对有现货库存的订单编制利库计划，减少重复生产'],
        ['计划管理', 'PMPM-A-03', '炼钢计划变更流程', '对已下发计划进行变更申请和审批'],
    ]
    make_table(doc, hdrs, data, [2.0, 2.5, 3.5, 8.0])

    # 3.3
    h(doc, '3.3  功能层级', 2)
    hdrs3 = ['子模块', '一级菜单', '二级菜单（页面名称）', '功能编码', '说明']
    data3 = [
        ['目标管理', '目标管理', 'BOM钢种合金原料配比', 'PMMB001', '维护钢种与合金原料的配比关系基础数据'],
        ['目标管理', '目标管理', '工作天数',             'PMMB002', '按月维护各工序的工作天数'],
        ['目标管理', '目标管理', '年度钢材生产计划',     'PMMB003', '制定年度钢材生产目标，按钢种分解'],
        ['目标管理', '目标管理', '月订单月原料推移',     'PMMB004', '展示月度订单与原料的推移分析看板'],
        ['订单管理', '订单管理', '炼钢订单',             'PMOM001', '管理从BIP同步的炼钢订单，支持查询和状态跟踪'],
        ['订单管理', '订单管理', '产能维护',             'PMOM002', '维护各轧线的月度产能约束数据'],
        ['订单管理', '订单管理', '交期评估',             'PMOM003', '对炼钢订单进行产能负荷评估并确认承诺交期'],
        ['计划管理', '计划管理', '炼钢计划编制',         'PMPM007', '编制、确认和下发炼钢生产计划'],
        ['计划管理', '计划管理', '利库计划',             'PMPM008', '编制针对库存现货的利库计划'],
        ['计划管理', '计划管理', '计划变更申请',         'PMPM009', '对已下发计划发起变更申请并流转审批'],
        ['计划管理', '计划管理', '计划完成率看板',       'PMPM010', '展示计划完成率多维统计分析（只读看板）'],
    ]
    make_table(doc, hdrs3, data3, [2.0, 2.0, 3.5, 2.0, 6.5])

# ─────────────────────────────────────────────────────────────
#  第 4.3 章：计划管理详细设计
# ─────────────────────────────────────────────────────────────
# IPO 五列表头和列宽（全局统一）
IPO_HDRS   = ['序号', '按钮/功能', '输入（Input）', '处理逻辑（Process）', '输出（Output）']
IPO_WIDTHS = [0.8, 2.2, 3.5, 6.2, 3.3]   # 合计 16.0 cm

def _sub_label(doc, text, indent_pt=21, sb=6, sa=2):
    """小节内标签段落（黑体 10.5pt 加粗）"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, cn_font=FONT_HEAD, size=10.5, bold=True)
    pf = p.paragraph_format
    pf.left_indent  = Pt(indent_pt)
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    return p

def build_ch43(doc):
    h(doc, '4.  详细设计', 1)
    h(doc, '4.3  计划管理', 2)

    # ── 4.3.1 系统流程清单 ────────────────────────────────────
    h(doc, '4.3.1  系统流程清单', 3)
    make_table(doc,
        ['流程编码', '流程名称', '主要涉及岗位', '流程简述'],
        [
            ['PMPM-A-01', '炼钢计划编制流程', '计划员 / 生产主管 / 炼钢工长',
             '以炼钢订单为输入，计划员编制炼钢计划草稿，经生产主管确认后下发至炼钢工序执行'],
            ['PMPM-A-02', '利库计划流程', '计划员 / 物料员 / 生产主管',
             '对可用库存现货满足的订单编制利库计划，经审批后下发，避免重复生产'],
            ['PMPM-A-03', '炼钢计划变更流程', '计划员 / 生产主管 / 调度员',
             '对已下发计划发起变更申请，经审批后更新计划并通知相关岗位'],
        ],
        [2.5, 3.0, 3.5, 7.0])

    # ── 4.3.2 系统流程说明 ────────────────────────────────────
    h(doc, '4.3.2  系统流程说明', 3)
    h(doc, '4.3.2.1  PMPM-A-01  炼钢计划编制流程', 4)
    _a01_png = generate_flow_png(FLOW_PMPM_A01, os.path.join(_IMG_DIR, 'PMPM-A-01.png'))
    generate_flow_drawio(FLOW_PMPM_A01, os.path.join(_IMG_DIR, 'PMPM-A-01.drawio'))
    insert_flow_img(doc, _a01_png, width_cm=16.0,
                    caption='图 4.3-1  炼钢计划编制流程（PMPM-A-01）— 泳道图')
    body(doc, '该流程以 BIP 同步的炼钢订单为输入，由计划员在炼钢计划编制页面完成计划数量、'
              '计划日期、所属轧线等字段的录入，形成草稿状态的炼钢计划；经生产主管确认后，'
              '状态流转为「已确认」，再由计划员下发，系统推送通知至炼钢工长，驱动工序执行。'
              '计划状态路径：草稿 → 已确认 → 已下发（单向不可逆；已下发需变更须走 PMPM-A-03 流程）。')

    _sub_label(doc, '（一）活动说明', indent_pt=0, sb=8, sa=4)
    make_table(doc,
        ['活动编号', '活动名称', '对应画面', '活动说明'],
        [
            ['PMPM-A-01-E-01', '炼钢订单接收', 'PMPM007 / 列表页',
             '系统从订单管理模块实时同步当期有效炼钢订单（状态=钢坯确认），无需手工导入；'
             '计划员可在列表页通过查询区筛选目标订单，确认本期待排产清单。'],
            ['PMPM-A-01-E-02', '炼钢计划编制', 'PMPM007 / 新增页',
             '计划员点击【新增】，选择炼钢订单后系统自动带出钢种、规格；计划员填写炼钢序号、'
             '计划数量（吨）、计划开始/结束日期、所属轧线后【确认】保存，生成草稿状态计划条目；'
             '草稿状态可反复修改，直至主管确认为止。'],
            ['PMPM-A-01-E-03', '计划确认', 'PMPM007 / 列表页',
             '生产主管在列表页选中一条或多条草稿状态计划行，点击【确认计划】；'
             '系统记录 confirm_user（操作人）和 confirm_time（系统时间），状态更新为「已确认」；'
             '已确认计划不可再修改或删除。'],
            ['PMPM-A-01-E-04', '计划下发', 'PMPM007 / 列表页',
             '计划员选中一条或多条已确认状态计划行，点击【下发计划】；系统更新状态为「已下发」，'
             '调用消息推送接口向炼钢工长及调度员发送计划到达通知；'
             '已下发计划全字段只读，如需变更须发起 PMPM-A-03 变更流程。'],
        ],
        [3.0, 2.5, 2.0, 8.5])

    # ── 4.3.3 系统流程与作业画面对照表 ────────────────────────
    h(doc, '4.3.3  系统流程与作业画面对照表', 3)
    make_table(doc,
        ['系统流程编号', '活动编号', '作业编号及名称', '页面说明'],
        [
            ['PMPM-A-01', 'PMPM-A-01-E-01', 'PMPM007  炼钢计划编制 — 列表页',
             '计划员在查询区筛选炼钢订单；列表页展示当期炼钢订单及已编制的计划条目'],
            ['PMPM-A-01', 'PMPM-A-01-E-02', 'PMPM007  炼钢计划编制 — 新增/修改弹窗',
             '点击【新增】弹出新增表单；选中草稿行点击【修改】弹出预填修改表单'],
            ['PMPM-A-01', 'PMPM-A-01-E-03', 'PMPM007  炼钢计划编制 — 列表页',
             '选中草稿状态计划行，点击【确认计划】，状态流转至「已确认」'],
            ['PMPM-A-01', 'PMPM-A-01-E-04', 'PMPM007  炼钢计划编制 — 列表页',
             '选中已确认状态计划行，点击【下发计划】，系统推送通知并更新状态至「已下发」'],
        ],
        [2.5, 3.0, 4.0, 6.5])

    # ── 4.3.4 功能设计 ────────────────────────────────────────
    h(doc, '4.3.4  功能设计', 3)

    # ── 4.3.4.3 炼钢订单（只读查询页） ──────────────────────
    h(doc, '4.3.4.3  【PMPM003】炼钢订单', 4)
    h(doc, '4.3.4.3.1  画面逻辑（原型）', 5)
    _ph003 = img_placeholder_png('PMPM003 炼钢订单\n界面原型图（列表查询页）',
                                  os.path.join(_IMG_DIR, 'PMPM003_prototype.png'))
    insert_flow_img(doc, _ph003, width_cm=14.0,
                    caption='图 4.3.4.3-1  PMPM003 炼钢订单 — 界面原型（待 UI 设计稿替换）')
    make_table(doc,
        ['区域', '内容说明', '交互/约束'],
        [
            ['入口路径', '计划管理 → 炼钢订单', '主菜单 → 二级菜单路径'],
            ['页面类型', '列表查询页（只读，无增删改）',
             '数据从订单管理模块（PMOM）实时同步，本页仅展示和导出'],
            ['查询条件区',
             '炼钢订单号（文本，模糊，选填）\n'
             '钢种（下拉，选填）\n'
             '规格（文本，模糊，选填）\n'
             '要求交期（日期范围，选填）\n'
             '状态（下拉，多选；选项：待处理 / 钢坯确认 / 已关闭）',
             '默认查全量，按要求交期升序；\n点击【重置】清空条件'],
            ['操作按钮区', '【搜索】【重置】　　（右对齐）【导出】',
             '本页面无新增/修改/删除按钮，全部只读'],
            ['列表展示区',
             '序号 / 炼钢订单号 / 钢种 / 规格 / 需求数量（吨）/\n'
             '已排产数量（吨）/ 未排产数量（吨，= 需求 - 已排产）/\n'
             '要求交期 / 状态（待处理=灰 / 钢坯确认=蓝 / 已关闭=绿）/\n'
             '来源系统（BIP / 手工）/ 同步时间',
             '未排产数量 < 0 时以红色字体显示（超排预警）；\n分页 20 条；列头支持排序'],
            ['说明',
             '本页为炼钢计划编制的数据来源参考页；\n'
             '计划员可在此确认各订单的未排产数量，再切换至 PMPM007 进行计划录入',
             '不支持在本页直接操作，需切换到计划编制页'],
        ],
        [2.0, 8.0, 6.0])

    h(doc, '4.3.4.3.2  处理逻辑（IPO）', 5)
    make_table(doc, IPO_HDRS,
        [
            ['1', '【搜索】',
             '炼钢订单号（可选，模糊）\n钢种（可选，精确）\n规格（可选，模糊）\n'
             '要求交期范围（可选）\n状态（可选，多选）',
             '按条件联合查询 pmpm_order 视图（源于 PMOM 模块同步）；\n'
             '「未排产数量」= 需求数量 - SUM(已有草稿/已确认/已下发计划数量)',
             'Toast「共 N 条记录」；列表刷新；\n'
             '无数据时显示「暂无数据」'],
            ['2', '【重置】', '/',
             '清空所有查询条件至默认值',
             '查询条件区恢复初始状态；\n列表不自动刷新'],
            ['3', '【导出】', '/',
             '将当前查询结果集（全量，不分页）生成 Excel；\n'
             '文件名：炼钢订单_YYYYMMDD_HHmmss.xlsx；\n'
             '列：订单号、钢种、规格、需求数量、已排产数量、未排产数量、要求交期、状态、来源系统、同步时间',
             '浏览器触发文件下载；\n查询结果为空时 Toast「无数据可导出」'],
        ],
        IPO_WIDTHS, font_size=9)

    # ── 4.3.4.7 炼钢计划编制 ─────────────────────────────────
    h(doc, '4.3.4.7  【PMPM007】炼钢计划编制', 4)

    # ── H5: 画面逻辑（原型） ──────────────────────────────────
    h(doc, '4.3.4.7.1  画面逻辑（原型）', 5)
    _ph007 = img_placeholder_png(
        'PMPM007 炼钢计划编制\n界面原型图（列表页 + 新增/修改弹窗）',
        os.path.join(_IMG_DIR, 'PMPM007_prototype.png'))
    insert_flow_img(doc, _ph007, width_cm=14.0,
                    caption='图 4.3.4.7-1  PMPM007 炼钢计划编制 — 界面原型（待 UI 设计稿替换）')
    make_table(doc,
        ['区域', '内容说明', '交互/约束'],
        [
            ['入口路径', '计划管理 → 炼钢计划编制', '主菜单 → 二级菜单路径'],
            ['页面类型', '列表页 + 新增/修改弹窗（同一路由）', '弹窗宽度建议 800px，高度自适应'],
            ['查询条件区',
             '炼钢订单号（文本输入，模糊匹配，选填）\n'
             '钢种（下拉单选，选填）\n'
             '计划日期范围（日期范围选择器，起止，选填）\n'
             '状态（下拉单选，选填；选项：草稿 / 已确认 / 已下发 / 全部）',
             '默认查全量，按计划开始日期降序；\n点击【重置】清空所有条件'],
            ['操作按钮区',
             '【搜索】【重置】　　【新增】【修改】【删除】【确认计划】【下发计划】　　（右对齐）【导出】',
             '修改/删除/确认/下发 在未选中行时置灰不可点'],
            ['列表展示区',
             '序号 / 炼钢订单号 / 炼钢序号 / 钢种 / 规格 / 计划数量（吨）/\n'
             '计划开始日期 / 计划结束日期 / 所属轧线 /\n'
             '状态（草稿=灰 / 已确认=蓝 / 已下发=绿）/ 创建人 / 创建时间 / 行内操作（修改/删除）',
             '分页：每页 20 条；支持按列头排序；复选框支持多选'],
            ['新增/修改弹窗',
             '炼钢订单号（搜索下拉，必填） → 自动带出钢种（只读）、规格（只读）\n'
             '炼钢序号（整数，必填，同订单下不可重复）\n'
             '计划数量（正数，精度 0.00，单位：吨，必填）\n'
             '计划开始日期（日期选择，格式 YYYY-MM-DD，必填）\n'
             '计划结束日期（日期选择，须 ≥ 开始日期，必填）\n'
             '所属轧线（下拉，来自轧线维护表，必填）\n'
             '备注（文本域，≤ 200 字，选填）',
             '弹窗底部：【确认】【取消】\n修改时炼钢订单号字段只读（不可更换关联订单）'],
            ['状态规则',
             '草稿：可修改 / 可删除 / 可确认\n已确认：不可修改 / 不可删除 / 可下发\n已下发：全字段只读，变更须走审批流程',
             '状态变更均需记录操作人和操作时间'],
        ],
        [2.0, 8.0, 6.0])

    # ── H5: 处理逻辑（IPO） ───────────────────────────────────
    h(doc, '4.3.4.7.2  处理逻辑（IPO）', 5)

    # IPO (1) 列表页
    _sub_label(doc, '（1）列表页', indent_pt=0, sb=4, sa=2)
    make_table(doc, IPO_HDRS,
        [
            ['1', '【搜索】',
             '炼钢订单号（可选，模糊）\n钢种（可选，精确）\n计划日期范围（可选）\n状态（可选，精确）',
             '按条件查询 pmpm_sm_plan 表（排除软删除记录）；\n无条件时查全量；按计划开始日期降序返回',
             'Toast「共 N 条记录」；\n列表刷新；\n无数据时显示「暂无数据」'],
            ['2', '【重置】',
             '/',
             '清空所有查询条件字段至默认值（空 / 全部）',
             '查询条件区恢复初始状态；\n列表不自动刷新（需手动点击【搜索】）'],
            ['3', '【新增】',
             '/',
             '校验当前用户角色有新增权限（角色：计划员）；\n弹出新增弹窗，所有字段初始化为空',
             '弹出【新增炼钢计划】弹窗'],
            ['4', '【修改】',
             '选中行（单选）\n状态须为「草稿」',
             '①校验恰好选中一行，否则提示「请选择一条记录」；\n'
             '②校验所选行状态为「草稿」，非草稿则 Toast「仅草稿状态可修改」；\n'
             '③按 plan_id 查询完整数据回填弹窗',
             '弹出【修改炼钢计划】弹窗，字段预填当前值；\n炼钢订单号字段只读'],
            ['5', '【删除】',
             '选中行（单选或多选）\n所选行状态须全为「草稿」',
             '①校验所选行状态全为「草稿」，有非草稿行则 Toast 提示并高亮非法行；\n'
             '②弹出二次确认：「确定删除选中的 X 条记录吗？」；\n'
             '③确认后软删除（is_deleted=1）并写入 delete_user、delete_time',
             '处理成功：Toast「已删除 X 条记录」，列表刷新；\n'
             '处理失败：Toast「删除失败，请重试」；\n未选中行时按钮置灰'],
            ['6', '【确认计划】',
             '选中行（单选或多选）\n所选行状态须全为「草稿」',
             '①校验所选行状态全为「草稿」；\n'
             '②批量更新 status = 已确认；\n'
             '③写入 confirm_user（当前操作人）、confirm_time（NOW()）；\n'
             '④写入操作日志',
             '处理成功：Toast「已确认 X 条计划」，列表刷新，状态列变为「已确认」（蓝色）；\n'
             '无草稿行可选时按钮置灰'],
            ['7', '【下发计划】',
             '选中行（单选或多选）\n所选行状态须全为「已确认」',
             '①校验所选行状态全为「已确认」；\n'
             '②批量更新 status = 已下发；\n'
             '③写入 issue_user、issue_time；\n'
             '④调用 API-PMPM-001 向炼钢工长/调度员推送通知；\n'
             '⑤写入操作日志',
             '处理成功：Toast「已下发 X 条计划」，列表刷新，状态列变为「已下发」（绿色）；\n'
             '推送失败时：Toast 警告「计划已下发，但通知推送失败，请手动告知相关人员」'],
            ['8', '【导出】',
             '/',
             '将当前查询条件下的全量结果集（排除软删除）生成 Excel；\n'
             '文件名：炼钢计划_YYYYMMDD_HHmmss.xlsx；\n'
             '列：炼钢订单号、炼钢序号、钢种、规格、计划数量（吨）、\n'
             '计划开始日期、计划结束日期、所属轧线、状态、创建人、创建时间',
             '浏览器触发文件下载；\n查询结果为空时 Toast「无数据可导出」'],
        ],
        IPO_WIDTHS, font_size=9)

    # IPO (2) 新增/修改弹窗
    _sub_label(doc, '（2）新增炼钢计划 / 修改炼钢计划 弹窗', indent_pt=0, sb=8, sa=2)
    make_table(doc, IPO_HDRS,
        [
            ['1', '页面加载\n字段初始化',
             '新增：所有字段为空\n修改：按 plan_id 回填已有数据',
             '字段显示规则：\n'
             '· 炼钢订单号：搜索下拉，输入关键字模糊匹配炼钢订单表中 status=钢坯确认 的记录；\n'
             '  修改时此字段只读（显示订单号文本，不可更改关联）\n'
             '· 钢种/规格：选中订单号后自动带出，只读展示\n'
             '· 炼钢序号：整数，新增时默认值=该订单已有计划数+1，可手动调整\n'
             '· 计划数量：正数，精度 2 位小数，单位：吨，前端即时校验 > 0\n'
             '· 计划开始日期：日期选择器，YYYY-MM-DD，新增时不可早于今天\n'
             '· 计划结束日期：日期选择器，须 ≥ 计划开始日期（前端即时联动校验）\n'
             '· 所属轧线：下拉单选，数据来自 base_rolling_line 表中 is_active=1 的记录\n'
             '· 备注：文本域，最大 200 字，选填',
             '订单号选中后钢种/规格即时联动带出；\n'
             '计划结束日期选择器禁用早于开始日期的日期'],
            ['2', '【确认】',
             '（无额外入参，\n读取弹窗各字段值）',
             '（一）前端实时校验：\n'
             '  ①所有必填项非空；\n'
             '  ②计划结束日期 ≥ 计划开始日期；\n'
             '  ③计划数量 > 0\n'
             '（二）后端业务校验：\n'
             '  ①炼钢订单号存在且状态有效；\n'
             '  ②同炼钢订单号 + 炼钢序号组合在计划表中不存在（新增时）\n'
             '（三）数据处理：\n'
             '  · 新增：INSERT pmpm_sm_plan，status=草稿，\n'
             '    create_user=当前用户，create_time=NOW()\n'
             '  · 修改：UPDATE pmpm_sm_plan SET ... WHERE id=? AND status=草稿\n'
             '（四）触发事件：无（仅下发时才推送消息）',
             '前端校验失败：在字段下方显示红色提示，不阻断点击但不提交；\n'
             '后端校验失败：弹窗顶部显示 Banner「炼钢序号已存在，请重新输入」；\n'
             '处理成功：Toast「保存成功」，关闭弹窗，列表刷新'],
            ['3', '【取消】',
             '/',
             '不执行任何数据写入，直接关闭弹窗',
             '关闭弹窗，列表页无任何变化'],
        ],
        IPO_WIDTHS, font_size=9)

    # ── 4.3.4.8 炼钢计划确认 ─────────────────────────────────
    h(doc, '4.3.4.8  【PMPM008】炼钢计划确认', 4)
    h(doc, '4.3.4.8.1  画面逻辑（原型）', 5)
    _ph008 = img_placeholder_png('PMPM008 炼钢计划确认\n界面原型图（待确认工作台）',
                                  os.path.join(_IMG_DIR, 'PMPM008_prototype.png'))
    insert_flow_img(doc, _ph008, width_cm=14.0,
                    caption='图 4.3.4.8-1  PMPM008 炼钢计划确认 — 界面原型（待 UI 设计稿替换）')
    make_table(doc,
        ['区域', '内容说明', '交互/约束'],
        [
            ['入口路径', '计划管理 → 炼钢计划确认', '主菜单 → 二级菜单路径'],
            ['页面类型', '列表页（待确认计划审核工作台）',
             '默认仅展示「草稿」状态计划，避免干扰已完成数据'],
            ['查询条件区',
             '炼钢订单号（文本，模糊，选填）\n'
             '钢种（下拉，选填）\n'
             '计划日期范围（选填）\n'
             '提交人（文本，选填）',
             '默认 status=草稿；切换「全部」可查历史已确认/退回记录'],
            ['操作按钮区',
             '【搜索】【重置】　　【确认选中】【退回选中】　　（右对齐）【导出】',
             '未选中行时【确认选中】【退回选中】置灰；\n仅车间主任角色可执行确认/退回操作'],
            ['列表展示区',
             '序号 / 炼钢订单号 / 炼钢序号 / 钢种 / 规格 /\n'
             '计划数量（吨）/ 计划开始日期 / 计划结束日期 / 所属轧线 /\n'
             '状态（草稿=灰 / 已确认=蓝 / 已退回=红）/\n'
             '提交人 / 提交时间 / 备注',
             '分页 20 条；复选框支持多选'],
            ['退回弹窗',
             '退回意见（文本域，必填，≤ 200 字）\n确定【退回】/【取消】',
             '退回意见不可为空；退回后状态变为「已退回」，计划员可重新修改提交'],
            ['状态规则',
             '草稿 → 已确认（操作：确认选中，操作人：车间主任）\n'
             '草稿 → 已退回（操作：退回选中，需填写退回意见）\n'
             '已退回 → 草稿（计划员修改后可重新提交）',
             '已确认状态不可退回，需走 PMPM-A-03 变更流程'],
        ],
        [2.0, 8.0, 6.0])

    h(doc, '4.3.4.8.2  处理逻辑（IPO）', 5)
    _sub_label(doc, '（1）列表页', indent_pt=0, sb=4, sa=2)
    make_table(doc, IPO_HDRS,
        [
            ['1', '页面加载\n默认查询',
             '/',
             '默认查询 pmpm_sm_plan 表 status=草稿 的记录（排除软删除）；\n'
             '按提交时间降序排列',
             '列表展示草稿状态计划；\n无数据时显示「暂无待确认的计划」'],
            ['2', '【搜索】',
             '炼钢订单号 / 钢种 / 计划日期 / 提交人（均可选）\n状态（默认=草稿）',
             '按条件查询 pmpm_sm_plan 表，支持 status 切换为全部',
             'Toast「共 N 条记录」；列表刷新'],
            ['3', '【确认选中】',
             '选中行（单选或多选）\n所选行 status 须全为「草稿」',
             '①校验选中行 status 全为「草稿」，有非草稿行 Toast「X 条记录状态不符，已跳过」；\n'
             '②弹出确认对话框：「确认以下 N 条炼钢计划？」（展示炼钢序号列表）；\n'
             '③点击【确定】：批量 UPDATE status=已确认，写入 confirm_user、confirm_time；\n'
             '④写入操作日志',
             '处理成功：Toast「已确认 N 条计划」，列表刷新；\n'
             '状态列变为「已确认」（蓝色）'],
            ['4', '【退回选中】',
             '选中行（单选或多选）\n所选行 status 须全为「草稿」\n退回意见（必填）',
             '①校验选中行 status 全为「草稿」；\n'
             '②弹出退回意见填写弹窗；\n'
             '③填写退回意见后点击【退回】：批量 UPDATE status=已退回，\n'
             '  写入 reject_user、reject_time、reject_reason；\n'
             '④系统发送站内通知至计划员',
             '处理成功：Toast「已退回 N 条计划」，列表刷新；\n'
             '状态列变为「已退回」（红色）；\n'
             '计划员收到通知「您有 N 条炼钢计划被退回，请重新修改」'],
        ],
        IPO_WIDTHS, font_size=9)

    # ── 4.3.4.9 炼钢计划下发 ─────────────────────────────────
    h(doc, '4.3.4.9  【PMPM009】炼钢计划下发', 4)
    h(doc, '4.3.4.9.1  画面逻辑（原型）', 5)
    _ph009 = img_placeholder_png('PMPM009 炼钢计划下发\n界面原型图（待下发工作台）',
                                  os.path.join(_IMG_DIR, 'PMPM009_prototype.png'))
    insert_flow_img(doc, _ph009, width_cm=14.0,
                    caption='图 4.3.4.9-1  PMPM009 炼钢计划下发 — 界面原型（待 UI 设计稿替换）')
    make_table(doc,
        ['区域', '内容说明', '交互/约束'],
        [
            ['入口路径', '计划管理 → 炼钢计划下发', '主菜单 → 二级菜单路径'],
            ['页面类型', '列表页（待下发工作台）',
             '默认仅展示「已确认」状态计划，突出待操作项'],
            ['查询条件区',
             '炼钢订单号（文本，模糊，选填）\n'
             '钢种（下拉，选填）\n'
             '确认日期范围（选填）\n'
             '确认人（文本，选填）',
             '默认 status=已确认；可切换为全部查已下发历史'],
            ['操作按钮区',
             '【搜索】【重置】　　【下发选中】　　（右对齐）【导出】【下发历史】',
             '未选中行时【下发选中】置灰；\n仅计划员角色可执行下发操作'],
            ['列表展示区',
             '序号 / 炼钢订单号 / 炼钢序号 / 钢种 / 规格 /\n'
             '计划数量（吨）/ 计划开始日期 / 计划结束日期 / 所属轧线 /\n'
             '状态（已确认=蓝 / 已下发=绿）/\n'
             '确认人 / 确认时间',
             '分页 20 条；复选框支持多选'],
            ['下发结果区',
             '下发完成后，在列表顶部显示 Banner：\n'
             '「已成功下发 N 条计划，通知已推送至 X 位相关人员」\n'
             '推送失败时：「已下发 N 条计划，但 M 条通知推送失败，请手动告知」',
             '推送失败不阻断业务，数据状态仍更新为「已下发」'],
            ['下发历史',
             '点击【下发历史】按钮，弹出历史记录面板：\n'
             '下发时间 / 下发人 / 下发数量 / 通知推送结果',
             '历史记录只读'],
        ],
        [2.0, 8.0, 6.0])

    h(doc, '4.3.4.9.2  处理逻辑（IPO）', 5)
    _sub_label(doc, '（1）列表页', indent_pt=0, sb=4, sa=2)
    make_table(doc, IPO_HDRS,
        [
            ['1', '页面加载\n默认查询',
             '/',
             '默认查询 pmpm_sm_plan 表 status=已确认 的记录（排除软删除）；\n'
             '按确认时间降序排列',
             '列表展示已确认待下发计划；\n无数据时显示「暂无待下发的计划」'],
            ['2', '【搜索】',
             '炼钢订单号 / 钢种 / 确认日期范围 / 确认人（均可选）\n状态（默认=已确认）',
             '按条件查询 pmpm_sm_plan 表',
             'Toast「共 N 条记录」；列表刷新'],
            ['3', '【下发选中】',
             '选中行（单选或多选）\n所选行 status 须全为「已确认」',
             '①校验选中行 status 全为「已确认」，有不符合行 Toast「已跳过」；\n'
             '②弹出确认对话框：「确认下发以下 N 条炼钢计划至 MES？」；\n'
             '③点击【确定】：\n'
             '  · 批量 UPDATE status=已下发，写入 issue_user、issue_time；\n'
             '  · 调用 API-PMPM-001（炼钢计划下发通知接口），传入 plan_id 列表；\n'
             '  · 写入操作日志，记录下发结果（推送成功/失败）',
             '处理成功：列表顶部 Banner「已成功下发 N 条计划，通知已推送」；\n'
             '状态列变为「已下发」（绿色）；\n'
             'API 推送失败时：Banner「已下发但通知推送失败」，\n'
             '失败详情写入日志，不回滚业务状态'],
            ['4', '【导出】', '/',
             '将当前查询结果集生成 Excel；\n'
             '文件名：炼钢计划下发记录_YYYYMMDD.xlsx；\n'
             '列：订单号、炼钢序号、钢种、规格、计划数量、日期、所属轧线、状态、下发人、下发时间',
             '浏览器触发文件下载'],
        ],
        IPO_WIDTHS, font_size=9)

    # ── 4.3.5 内部接口（API） ─────────────────────────────────
    h(doc, '4.3.5  内部接口（API）', 3)
    make_table(doc,
        ['序号', '接口编码', '接口名称', '来源模块', '去向模块', '接口说明（输入 / 输出 / 触发时机）'],
        [
            ['1', 'API-PMPM-001', '炼钢计划下发通知',
             'PMPM 计划管理',
             '消息服务 /\nMES 炼钢接口',
             '触发：计划状态变更为「已下发」时；\n'
             '输入：plan_id 列表、炼钢订单号、炼钢序号、计划数量、计划开始/结束日期、所属轧线；\n'
             '输出：推送结果码（200=成功，非200=失败）；\n'
             '接收方：炼钢工长账号（App 推送 + 系统内消息），调度员账号（系统内消息）'],
            ['2', 'API-PMPM-002', '炼钢订单数据查询',
             'PMOM 订单管理',
             'PMPM 计划管理',
             '触发：新增弹窗中炼钢订单号字段输入关键字时；\n'
             '输入：keyword（订单号关键字，模糊）、status（=钢坯确认）；\n'
             '输出：炼钢订单列表（订单号、钢种、规格、需求数量、要求交期）；\n'
             '分页：前 20 条匹配结果'],
        ],
        [0.6, 2.2, 2.5, 2.0, 2.2, 6.5])

# ─────────────────────────────────────────────────────────────
#  主流程
# ─────────────────────────────────────────────────────────────
def main():
    # ── 以参考模板为基础，继承全部样式定义 ────────────────────
    doc = Document(TPLPATH)

    # 清空正文内容，保留样式 / 编号清单 / sectPr
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    # 移除 Heading 1-4 样式的自动编号（我们手动写编号文字）
    for lv in range(1, 5):
        try:
            st = doc.styles[f'Heading {lv}']
            pPr = st.element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)
        except Exception:
            pass

    # ── 页面设置：A4 竖向，匹配模板 Section 0 页边距 ──────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.27)
    # 去除模板的 landscape 标记，确保竖向
    pgSz = section._sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.attrib.pop(qn('w:orient'), None)

    # ── 页眉 / 页脚 ─────────────────────────────────────────
    setup_header(section)
    setup_footer(section)

    # ── 固定页：封面 1 / 封面 2 / 修订记录 / 目录 ───────────
    build_cover1(doc)
    build_cover2(doc)
    build_revision(doc)
    build_toc(doc)

    # ── 正文 ────────────────────────────────────────────────
    build_ch1(doc)
    build_ch2(doc)
    build_ch3(doc)
    build_ch43(doc)

    # ── 强制 Word 打开时自动更新所有域（包括目录）────────────
    update_flds = OxmlElement('w:updateFields')
    update_flds.set(qn('w:val'), 'true')
    doc.settings.element.append(update_flds)

    # ── 保存 ────────────────────────────────────────────────
    out = os.path.join(_OUT_DIR, '华新计划模块需求设计说明书_v2.0.docx')
    doc.save(out)
    print(f'✅  已生成：{out}')

if __name__ == '__main__':
    main()
